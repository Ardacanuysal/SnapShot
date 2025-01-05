from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('SnapShot Proje Dokümantasyonu', 0)

# Project Description
doc.add_heading('Proje Tanımı:', level=1)
doc.add_paragraph(
    "SnapShot, görüntü işleme teknikleri kullanarak yüz hatlarını belirli matris metrikleriyle birleştirip "
    "yüze entegre efekt uygulayan bir web platformudur. Proje, kullanıcının cihazındaki kamera iznini talep eder "
    "ve bu canlı görüntü üzerinden çeşitli görsel efektler uygular. Kullanıcılar, basit bir kullanıcı arayüzü (UI/UX) "
    "ile efektleri seçebilir ve kamerada anlık olarak bu efektleri görebilir.\n\n"
    "Teknolojiler:\n"
    "- Backend: Flask\n"
    "- Frontend: HTML, CSS, React Native\n"
    "- Assets: PNG dosyaları kullanılarak görseller korunmaktadır.\n"
    "- Görsel İşlemeler: Kontrast, zıtlık gibi görsel ayarların yapılabilmesi.\n"
    "- Kamera Entegrasyonu: Canlı kamera görüntüsü üzerinden API ile veri gönderme işlemi Flask ile gerçekleştirilmiştir.\n"
    "- UI/UX Tasarımı: Kullanıcı arayüzü sade ve basit tutulmuş olup, butonlarla kullanıcı kendi efektini seçebilmekte "
    "ve anlık olarak yüzüne efekt uygulamaktadır.\n"
)

# Features Section
doc.add_heading('Proje Özellikleri (Features):', level=1)
doc.add_paragraph(
    "- **Gerçek Zamanlı Video Akışı:** Yüz tespiti ile canlı video akışı sağlanır.\n"
    "- **Efekt Uygulama:** Kullanıcılar, gözlük, kalp emojisi, köpek görseli gibi efektleri seçebilir ve kamerada anlık olarak yüzlerine entegre edebilir.\n"
    "- **Video Efekt Kontrolü:** Sıcaklık (temperature), siyah beyaz (black & white), kontrast gibi görsel işlemler yapılabilir.\n"
    "- **'Dog Encoding' Aktif Etme:** Köpek efekti, etkinleştirildiğinde bir mesaj görüntülenir ve kullanıcıya geri bildirim sağlanır.\n"
)

# Installation Steps
doc.add_heading('Kurulum Adımları (Installation Steps):', level=1)
doc.add_paragraph(
    "1. Repository'i klonlayın.\n"
    "2. Proje klasörüne gidin.\n"
    "3. Sanal ortam oluşturun ve etkinleştirin.\n"
    "4. Bağımlılıkları yüklemek için `pip install -r requirements.txt` komutunu çalıştırın.\n"
    "5. Flask uygulamasını başlatmak için `python app.py` komutunu kullanın.\n"
)

# API Endpoints Section
doc.add_heading('API Endpoints:', level=1)
doc.add_paragraph(
    "- `/`: Ana sayfayı, video akışını gösterir.\n"
    "- `/video_feed`: Yüz tespiti ve efektlerle video akışını sağlar.\n"
    "- `/select_image (POST)`: Uygulamak istediğiniz görsel efekti seçer (gözlük, kalp, emoji, köpek).\n"
    "- `/set_contrast (POST)`: Video akışının kontrast seviyesini ayarlar.\n"
    "- `/start_dog_encoding (POST)`: Köpek efekti aktif hale gelir.\n"
    "- `/stop_dog_encoding (POST)`: Köpek efekti devre dışı bırakılır.\n"
)

# Example API Calls Section
doc.add_heading('API Örnek Çağrıları (Example API Calls):', level=1)
doc.add_paragraph(
    "- **Kontrast Seviyesi Ayarlama:**\n"
    "  ```bash\n"
    "  curl -X POST http://127.0.0.1:5000/set_contrast -d '{\"contrast\": 70}' -H \"Content-Type: application/json\"\n"
    "  ```\n"
    "- **Köpek Encoding Etkinleştirme:**\n"
    "  ```bash\n"
    "  curl -X POST http://127.0.0.1:5000/start_dog_encoding\n"
    "  ```\n"
    "- **Köpek Encoding Devre Dışı Bırakma:**\n"
    "  ```bash\n"
    "  curl -X POST http://127.0.0.1:5000/stop_dog_encoding\n"
    "  ```\n"
)

# Save the document to a file
doc_path = "/mnt/data/SnapShot_Proje_Dokumantasyonu.docx"
doc.save(doc_path)

doc_path
